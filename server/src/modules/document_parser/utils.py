import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from lxml import etree
from bs4 import BeautifulSoup
import tiktoken
import re
from typing import Dict, List, Any
import logging

class DocumentProcessor:
    """Main document processing class that handles multiple file types"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.pdf_methods = {
            'pymupdf': self._process_with_pymupdf,
            'pypdf2': self._process_with_pypdf2
        }
    
    def process_pdf(self, file_path: str, method: str = 'pymupdf') -> Dict[str, Any]:
        """Process PDF files with multiple fallback methods"""
        try:
            return self.pdf_methods[method](file_path)
        except Exception as e:
            self.logger.warning(f"Method {method} failed: {e}. Trying fallback...")
            # Try the other method as fallback
            fallback_method = 'pypdf2' if method == 'pymupdf' else 'pymupdf'
            return self.pdf_methods[fallback_method](file_path)
    
    def _process_with_pymupdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF using PyMuPDF - better for complex PDFs"""
        doc = fitz.open(file_path)
        content = {
            'text': '',
            'pages': [],
            'metadata': {}
        }
        
        # Get document metadata
        content['metadata'] = doc.metadata
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            content['text'] += page_text + '\n'
            content['pages'].append({
                'page_number': page_num + 1,
                'text': page_text
            })
        
        doc.close()
        return content
    
    def _process_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Process PDF using PyPDF2 - lighter weight option"""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            content = {
                'text': '',
                'pages': [],
                'metadata': pdf_reader.metadata or {}
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                content['text'] += page_text + '\n'
                content['pages'].append({
                    'page_number': page_num + 1,
                    'text': page_text
                })
            
            return content
    
    def process_word_doc(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents (.docx only for now)"""
        if not file_path.endswith('.docx'):
            raise ValueError("Only .docx files are supported currently")
        
        doc = Document(file_path)
        content = {
            'text': '',
            'paragraphs': [],
            'tables': []
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })
                content['text'] += para.text + '\n'
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'table_index': table_idx,
                'data': table_data
            })
            
            # Add table text to main content
            for row in table_data:
                content['text'] += ' | '.join(row) + '\n'
        
        return content
    
    def process_xml(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            tree = etree.parse(file_path)
            root = tree.getroot()
            
            return {
                'text': self._extract_text_from_xml(root),
                'structure': self._xml_to_dict(root),
                'metadata': {
                    'root_tag': root.tag,
                    'namespace': root.nsmap if hasattr(root, 'nsmap') else {}
                }
            }
        except Exception as e:
            # Fallback to BeautifulSoup for malformed XML
            self.logger.warning(f"lxml failed, trying BeautifulSoup: {e}")
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'xml')
                return {
                    'text': soup.get_text(),
                    'structure': str(soup.prettify()[:1000]) + "..." if len(str(soup.prettify())) > 1000 else str(soup.prettify())
                }
    
    def _extract_text_from_xml(self, element) -> str:
        """Extract all text content from XML element"""
        texts = []
        if element.text and element.text.strip():
            texts.append(element.text.strip())
        
        for child in element:
            texts.append(self._extract_text_from_xml(child))
        
        return ' '.join(filter(None, texts))
    
    def _xml_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary"""
        result = {}
        
        if element.text and element.text.strip():
            result['_text'] = element.text.strip()
        
        if element.attrib:
            result['_attributes'] = element.attrib
        
        for child in element:
            child_data = self._xml_to_dict(child)
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
        
        return result


class SmartTextChunker:
    """Intelligent text chunking that respects sentence boundaries"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200, model_name: str = "cl100k_base"):
        self.chunk_size = chunk_size
        self.overlap = overlap
        try:
            self.encoder = tiktoken.get_encoding(model_name)
        except Exception:
            # Fallback: simple character-based counting
            self.encoder = None
        self.logger = logging.getLogger(__name__)
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Smart chunking that respects sentence boundaries"""
        if not text or not text.strip():
            return []
        
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = self._get_token_count(sentence)
            
            # If adding this sentence would exceed chunk size
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk from current sentences
                chunk_text = ' '.join(current_chunk)
                chunks.append(self._create_chunk(chunk_text, len(chunks), metadata))
                
                # Start new chunk with overlap
                overlap_sentences = self._get_overlap_sentences(current_chunk)
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(self._get_token_count(s) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(self._create_chunk(chunk_text, len(chunks), metadata))
        
        return chunks
    
    def _get_token_count(self, text: str) -> int:
        """Get token count for text"""
        if self.encoder:
            return len(self.encoder.encode(text))
        else:
            # Rough approximation: 1 token ≈ 4 characters
            return len(text) // 4
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting using regex
        # This can be improved with spaCy or NLTK for better accuracy
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Clean up sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Ignore very short fragments
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _get_overlap_sentences(self, sentences: List[str]) -> List[str]:
        """Get sentences for overlap based on token count"""
        overlap_sentences = []
        overlap_length = 0
        
        for sentence in reversed(sentences):
            sentence_length = self._get_token_count(sentence)
            if overlap_length + sentence_length <= self.overlap:
                overlap_sentences.insert(0, sentence)
                overlap_length += sentence_length
            else:
                break
        
        return overlap_sentences
    
    def _create_chunk(self, text: str, chunk_index: int, metadata: Dict) -> Dict:
        """Create a chunk object with metadata"""
        chunk_metadata = metadata.copy() if metadata else {}
        chunk_metadata.update({
            'chunk_index': chunk_index,
            'chunk_length': len(text),
            'token_count': self._get_token_count(text)
        })
        
        return {
            'text': text,
            'metadata': chunk_metadata
        }