from pathlib import Path
import fitz  # PyMuPDF for PDF
import docx
import xml.etree.ElementTree as ET

from .models import ParsedDocument
from .utils import clean_text, chunk_text


class DocumentParser:

    def parse(self, file_path: str) -> ParsedDocument:
        """Route file to correct parser based on extension."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            text = self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            text = self._parse_docx(file_path)
        elif ext == ".xml":
            text = self._parse_xml(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        clean = clean_text(text)
        chunks = chunk_text(clean)

        return ParsedDocument(
            source=os.path.basename(file_path),
            content=chunks,
            metadata={"file_type": ext}
        )

    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text("text") + "\n"
        return text

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from Word docx."""
        doc = docx.Document(file_path)
        text = " ".join([para.text for para in doc.paragraphs])
        return text

    def _parse_xml(self, file_path: str) -> str:
        """Extract text from XML nodes."""
        tree = ET.parse(file_path)
        root = tree.getroot()
        text = " ".join([elem.text.strip() for elem in root.iter() if elem.text])
        return text
